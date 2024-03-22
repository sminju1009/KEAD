from konlpy.tag import Okt
from konlpy.tag import Kkma
from kiwipiepy import Kiwi
from collections import defaultdict
from docx import Document

okt = Okt()
kkma = Kkma()
kiwi = Kiwi()

# .docx 파일 열기
docx_path = '홍길동전.docx'
doc = Document(docx_path)

# docx 파일의 모든 텍스트를 읽어서 text 변수에 저장
text = ''
for paragraph in doc.paragraphs:
    text += paragraph.text + '\n'

# Kiwi를 사용하여 형태소 분석
tokens = kiwi.tokenize(text)

# 형태소 등장 횟수를 저장할 딕셔너리 생성
token_count = defaultdict(int)

# 형태소 분석 결과를 문서에 저장
result_doc = Document()
# result_doc.add_heading('형태소 분석 결과', level=1)

for token in tokens:
    # 명사(NNG)와 고유명사(NNP)인 경우에만 등장 횟수 저장
    if token[1] == 'NNG' or token[1] == 'NNP':
        # result_doc.add_paragraph(f'{token[0]}: {token[1]}')
        token_count[token[0]] += 1

# 각 형태소의 등장 횟수 출력
result_doc.add_heading('형태소 등장 횟수', level=1)
for token, count in token_count.items():
    result_doc.add_paragraph(f'{token}: {count}')

# 결과를 새로운 문서로 저장
result_doc.save('홍길동전_형태소_분석_결과_카운팅포함.docx')