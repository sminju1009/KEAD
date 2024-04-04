from konlpy.tag import Okt
from konlpy.tag import Kkma
from kiwipiepy import Kiwi
from docx import Document

okt = Okt()
kkma = Kkma()
kiwi = Kiwi()

# .docx 파일 열기
docx_path = '타임 머신 The Time Machine.docx'
doc = Document(docx_path)

# docx 파일의 모든 텍스트를 읽어서 text 변수에 저장
text = ''
for paragraph in doc.paragraphs:
    text += paragraph.text + '\n'

# Kiwi를 사용하여 형태소 분석
tokens = kiwi.tokenize(text)

# 형태소 분석 결과를 문서에 저장
result_doc = Document()
result_doc.add_heading('형태소 분석 결과', level=1)

for token in tokens:
    if token[1] == 'NNG':
        result_doc.add_paragraph(f'{token[0]}: {token[1]}')
        print('작업중')
    elif token[1] == 'NNP':
        doc.add_paragraph(f'{token[0]}: {token[1]}')
        print('작업중')

# token[0] 단어본문
# token[1] 품사구분
# token[2] 시작 글자수
# token[3] 단어 하나당 글자수        
    

# 결과를 새로운 문서로 저장
result_doc.save('타임 머신 The Time Machine_형태소_분석_결과.docx')


# print('OKT 형태소 분석 :',okt.morphs(text))
# print()
# print('OKT 품사 태깅 :',okt.pos(text))
# print()
# print('OKT 명사 추출 :',okt.nouns(text))
# print('**************************************************')
# print('꼬꼬마 형태소 분석 :',kkma.morphs(text))
# print()
# print('꼬꼬마 품사 태깅 :',kkma.pos(text))
# print()
# print('꼬꼬마 명사 추출 :',kkma.nouns(text))  

