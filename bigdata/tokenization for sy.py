from kiwipiepy import Kiwi
from docx import Document

kiwi = Kiwi()

# .docx 파일 열기
docx_path = '홍길동전.docx'
doc = Document(docx_path)

# docx 파일의 모든 텍스트를 읽어서 text 변수에 저장
text = ''
for paragraph in doc.paragraphs:
    text += paragraph.text + '\n'

# Kiwi를 사용하여 형태소 분석
tokens = kiwi.tokenize(text)

# 형태소 분석 결과를 문서에 저장
result_doc = Document()

# 토큰들을 띄어쓰기 한 칸으로 구분하여 문서에 추가
token_text = ' '.join([f'{token[0]}' for token in tokens if token[1] in ('NNG', 'NNP')])
result_doc.add_paragraph(token_text)

# token[0] 단어본문
# token[1] 품사구분
# token[2] 시작 글자수
# token[3] 단어 하나당 글자수        
    

# 결과를 새로운 문서로 저장
result_doc.save('홍길동전 for sy.docx')




